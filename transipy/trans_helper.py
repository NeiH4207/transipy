import ast
import docx
from time import sleep
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, List, Optional

import pandas as pd
import requests
from googletrans import Translator
from loguru import logger
from tqdm import tqdm

from transipy.utils import md5hash, split_df_by_group

# Constants
TRANSLATE_URL = "https://translate.googleapis.com/translate_a/single"
MAX_TEXT_LENGTH = 50000
CHUNK_SIZE = 5000

# Global fallback translator instance
global_translator = Translator()


def translate(text: str, src: str, dest: str) -> str:
    """
    Translate text from source language to destination language.
    
    If text is not a string or exceeds MAX_TEXT_LENGTH,
    the original text is returned.
    """
    if not isinstance(text, str):
        return text

    if len(text) > MAX_TEXT_LENGTH:
        logger.warning(
            f"Text length is too long: {len(text)} (max: {MAX_TEXT_LENGTH} characters), skipping..."
        )
        return text

    chunks = split_text_into_chunks(text)
    translated_chunks = translate_chunks(chunks, src, dest)
    return " ".join(translated_chunks)


def split_text_into_chunks(text: str) -> List[str]:
    """
    Split text into chunks not exceeding CHUNK_SIZE.
    Attempts to split on a period after CHUNK_SIZE to preserve sentence boundaries.
    """
    if len(text) <= CHUNK_SIZE:
        return [text]

    chunks = []
    start_point = 0
    while start_point < len(text):
        # Find a period after CHUNK_SIZE to break at sentence end.
        end_point = text.find(".", start_point + CHUNK_SIZE)
        if end_point == -1:
            end_point = len(text)
        chunk = text[start_point:end_point].strip()
        if chunk:
            chunks.append(chunk)
        start_point = end_point + 1  # Skip the period.
    return chunks


def translate_chunks(chunks: List[str], src: str, dest: str) -> List[str]:
    """
    Translate a list of text chunks. In case of an error,
    the original chunk is returned.
    """
    translated_chunks = []
    for chunk_text in chunks:
        try:
            translated = translate_single_chunk(chunk_text, src, dest)
            translated_chunks.append(translated)
        except Exception as exc:
            logger.exception(f"Error translating chunk: {exc}")
            translated_chunks.append(chunk_text)
    return translated_chunks


def translate_single_chunk(chunk_text: str, src: str, dest: str) -> str:
    """
    Translate a single text chunk using the Google Translate API.
    If the request fails, falls back to the googletrans library.
    """
    params = {
        "client": "gtx",
        "sl": src,
        "tl": dest,
        "dt": "t",
        "q": chunk_text,
    }
    try:
        response = requests.get(TRANSLATE_URL, params=params)
        response.raise_for_status()
        # Adjust the non-standard JSON response for literal_eval.
        fixed_text = response.text.replace("null", '"null"').replace("true", "True")
        translation_data = ast.literal_eval(fixed_text)
        return "".join(segment[0] for segment in translation_data[0])
    except requests.RequestException:
        logger.warning("Google Translate API request failed, falling back to googletrans library")
        return global_translator.translate(chunk_text, src=src, dest=dest).text


def translate_series(series: pd.Series, src: str = "en", dest: str = "vi") -> pd.Series:
    """
    Translate a pandas Series element-wise.
    """
    return series.apply(lambda text: translate(text, src, dest))


def translate_column(
    df: pd.DataFrame,
    column: str,
    src: str = "en",
    dest: str = "vi",
    chunk_size: int = 4,
    dictionary: Optional[Dict[str, str]] = None,
) -> pd.DataFrame:
    """
    Translate a specific column in a DataFrame, with caching via a dictionary.
    """
    dictionary = dictionary or {}
    df[column] = df[column].fillna("").astype(str).replace("nan", "")
    unique_values = df[column].unique()
    unique_series = pd.Series(unique_values, name=column)

    # Build a caching dictionary using MD5 hashes of lowercase text.
    caching_dict = {md5hash(key.lower()): value for key, value in dictionary.items()}

    # Split unique values into chunks.
    n_chunks = (len(unique_series) // chunk_size) + 1
    chunks = split_df_by_group(unique_series, n_chunks)

    with ThreadPoolExecutor() as executor:
        translated_results = list(
            tqdm(
                executor.map(lambda s: translate_series(s, src, dest), chunks),
                total=len(chunks),
                desc="Translating unique values",
            )
        )

    # Update caching dictionary with translated results.
    for chunk, result in zip(chunks, translated_results):
        for original, translated in zip(chunk, result):
            key_hash = md5hash(original.lower())
            if key_hash not in caching_dict:
                caching_dict[key_hash] = translated

    # Replace column values with their translations using the caching dictionary.
    df[column] = df[column].apply(lambda x: caching_dict.get(md5hash(x.lower()), x))
    df[column] = df[column].replace({"null": ""}).fillna("")
    sleep(0.1)
    return df


def translate_df(
    df: pd.DataFrame,
    columns: Optional[List[str]] = None,
    src: str = "en",
    dest: str = "vi",
    chunk_size: int = 4,
    dictionary: Optional[Dict[str, str]] = None,
    output_file: Optional[str] = None,
) -> pd.DataFrame:
    """
    Translate specified columns in a DataFrame.
    """
    translated_df = df.copy()
    dictionary = dictionary or {}

    if not columns:
        logger.warning("No columns specified for translation.")
        return translated_df

    for column in columns:
        if column not in translated_df.columns:
            logger.warning(f"Column '{column}' not found in DataFrame, skipping translation for this column.")
            continue

        if translated_df[column].dtype == "object":
            logger.info(f"Translating column: {column}")
            translate_column(translated_df, column, src, dest, chunk_size, dictionary)

    if output_file:
        translated_df.to_csv(output_file, index=False)

    return translated_df


def translate_excel(
    dfs: Dict[str, pd.DataFrame],
    sheets: Optional[List[str]] = None,
    src: str = "en",
    dest: str = "vi",
    chunk_size: int = 4,
    dictionary: Optional[Dict[str, str]] = None,
    output_file: Optional[str] = None,
) -> Dict[str, pd.DataFrame]:
    """
    Translate specified sheets in an Excel file represented as a dictionary of DataFrames.
    """
    dictionary = dictionary or {}
    sheets = sheets or list(dfs.keys())

    for sheet in tqdm(sheets, desc="Translating Excel sheets"):
        if sheet not in dfs:
            logger.warning(f"Sheet '{sheet}' not found in provided data, skipping...")
            continue

        logger.info(f"Translating sheet: {sheet}")
        df_sheet = dfs[sheet]
        dfs[sheet] = translate_df(
            df=df_sheet,
            columns=list(df_sheet.columns),
            src=src,
            dest=dest,
            chunk_size=chunk_size,
            dictionary=dictionary,
        )

    if output_file:
        with pd.ExcelWriter(output_file) as writer:
            for sheet in sheets:
                if sheet in dfs:
                    dfs[sheet].to_excel(writer, sheet_name=sheet, index=False)

    return dfs


def translate_docx(
    file_path: str,
    src: str = "en",
    dest: str = "vi",
    chunk_size: int = 4,
    dictionary: Optional[Dict[str, str]] = None,
    output_file: Optional[str] = None,
) -> docx.Document:
    """
    Translate a Word document (docx) and return a new translated Document.
    """
    dictionary = dictionary or {}
    original_doc = docx.Document(file_path)
    translated_doc = docx.Document()

    paragraphs = [para.text for para in original_doc.paragraphs]
    df = pd.DataFrame(paragraphs, columns=["text"])

    translated_df = translate_df(
        df=df,
        columns=["text"],
        src=src,
        dest=dest,
        chunk_size=chunk_size,
        dictionary=dictionary,
    )

    for text in translated_df["text"]:
        translated_doc.add_paragraph(text)

    if output_file:
        translated_doc.save(output_file)

    return translated_doc
